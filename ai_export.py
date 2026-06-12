#!/usr/bin/env python3
"""
AI Weekly Report Export Engine
导出引擎 — PDF/Word 一键导出
================================
用法:
    python ai_export.py                         # 导出HTML为Word+PDF
    python ai_export.py --format pdf             # 仅PDF
    python ai_export.py --format docx            # 仅Word
    python ai_export.py --html report.html       # 指定HTML文件
    python ai_export.py --watch                  # 监听模式(HTML变更自动导出)
"""

import json
import sys
import os
import argparse
import time
from pathlib import Path
from datetime import datetime
from typing import Optional


# ============================================================
# 配置
# ============================================================
CONFIG = {
    "default_html": "ai_weekly_report.html",
    "output_dir": "exports",
    "company_name": "AI Weekly Report",
    "author": "AI Research Team",
}


# ============================================================
# DOCX 导出 (python-docx)
# ============================================================

def export_to_docx(html_path: str, output_path: str = None) -> str:
    """
    将 HTML 周报转换为 Word 文档
    提取 HTML 中的文本内容，按标题层级排版
    """
    from html.parser import HTMLParser

    class DocxHTMLParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.sections = []       # [(level, title, content), ...]
            self.current_tag = None
            self.current_text = ""
            self.in_title = False
            self.in_skip = False     # 跳过 script/style
            self.skip_tags = {"script", "style", "noscript", "svg"}
            self.h_tags = {"h1", "h2", "h3", "h4"}

        def handle_starttag(self, tag, attrs):
            tag_l = tag.lower()
            if tag_l in self.skip_tags:
                self.in_skip = True
                return
            if tag_l in self.h_tags:
                self.current_tag = tag_l
                self.in_title = True
                self.current_text = ""

        def handle_endtag(self, tag):
            tag_l = tag.lower()
            if tag_l in self.skip_tags:
                self.in_skip = False
                return
            if tag_l in self.h_tags and self.in_title:
                level = int(tag_l[1])
                self.sections.append((level, self.current_text.strip(), []))
                self.current_tag = None
                self.in_title = False
                self.current_text = ""
            elif tag_l in ("p", "li", "div"):
                text = self.current_text.strip()
                if text and len(text) > 10 and self.sections:
                    self.sections[-1][2].append(text)
                self.current_text = ""

        def handle_data(self, data):
            if self.in_skip:
                return
            self.current_text += data

    # Parse HTML
    html_content = Path(html_path).read_text(encoding="utf-8")
    parser = DocxHTMLParser()
    parser.feed(html_content)

    # Build DOCX
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("⚠ python-docx 未安装。运行: pip install python-docx")
        return ""

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题页
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("AI 行业周报")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(59, 130, 246)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_paragraph()  # spacer

    # 正文
    for level, title, contents in parser.sections:
        if not title or len(title) < 3:
            continue

        heading = doc.add_heading(title, level=min(level, 3))
        for content in contents[:5]:  # 每节最多5段
            p = doc.add_paragraph(content[:300])
            p.style.font.size = Pt(10.5)

    # 页脚
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("— 内容由AI自动生成，仅供参考 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)
    run.font.italic = True

    # 保存
    if output_path is None:
        output_path = str(Path(html_path).parent / CONFIG["output_dir"] / 
                         f"AI_Weekly_Report_{datetime.now().strftime('%Y%m%d')}.docx")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Word 文档已导出: {output_path}")
    return output_path


# ============================================================
# PDF 导出
# ============================================================

def export_to_pdf(html_path: str, output_path: str = None) -> str:
    """
    将 HTML 周报转换为 PDF
    方案1: weasyprint (推荐)
    方案2: pdfkit (需要 wkhtmltopdf)
    方案3: 浏览器 headless (需要 playwright)
    """
    if output_path is None:
        output_path = str(Path(html_path).parent / CONFIG["output_dir"] /
                         f"AI_Weekly_Report_{datetime.now().strftime('%Y%m%d')}.pdf")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # 尝试 weasyprint
    try:
        from weasyprint import HTML
        HTML(filename=html_path).write_pdf(output_path)
        print(f"✅ PDF 已导出 (weasyprint): {output_path}")
        return output_path
    except ImportError:
        pass

    # 尝试 pdfkit
    try:
        import pdfkit
        options = {
            'page-size': 'A4',
            'margin-top': '15mm',
            'margin-bottom': '15mm',
            'margin-left': '10mm',
            'margin-right': '10mm',
            'encoding': 'UTF-8',
            'no-outline': None,
            'enable-local-file-access': None,
        }
        pdfkit.from_file(html_path, output_path, options=options)
        print(f"✅ PDF 已导出 (pdfkit): {output_path}")
        return output_path
    except ImportError:
        pass

    # 方案3: 浏览器 headless
    try:
        from playwright.sync_api import sync_playwright
        abs_path = f"file:///{Path(html_path).absolute().as_posix()}"
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(abs_path, wait_until="networkidle")
            page.pdf(path=output_path, format="A4", print_background=True)
            browser.close()
        print(f"✅ PDF 已导出 (playwright): {output_path}")
        return output_path
    except ImportError:
        pass

    # 方案4: 纯文本fallback
    print("⚠ 无可用的PDF引擎。安装 weasyprint: pip install weasyprint")
    print("  或 pdfkit: pip install pdfkit + 安装 wkhtmltopdf")
    print("  或 playwright: pip install playwright && playwright install chromium")
    return ""


def export_json_report(html_path: str, output_path: str = None) -> str:
    """导出JSON格式数据报告 (供程序化消费)"""
    if output_path is None:
        output_path = str(Path(html_path).parent / CONFIG["output_dir"] /
                         f"AI_Weekly_Data_{datetime.now().strftime('%Y%m%d')}.json")

    # 从HTML提取数据 (简化版)
    summary = {
        "title": "AI行业周报",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "format_version": "2.0",
        "export_engine": "ai_export.py",
    }
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"✅ JSON 数据已导出: {output_path}")
    return output_path


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="AI Weekly Report Export Engine")
    parser.add_argument("--html", default=None, help="HTML源文件路径")
    parser.add_argument("--output", default=None, help="输出目录/文件")
    parser.add_argument("--format", choices=["all", "pdf", "docx", "json"], default="all",
                       help="导出格式 (默认: all)")
    parser.add_argument("--watch", action="store_true", help="监听模式")
    args = parser.parse_args()

    html_file = args.html or os.path.join(os.path.dirname(__file__), CONFIG["default_html"])

    if not Path(html_file).exists():
        print(f"❌ HTML 文件不存在: {html_file}")
        sys.exit(1)

    def do_export():
        out_dir = args.output or str(Path(html_file).parent / CONFIG["output_dir"])
        fmt = args.format

        results = {}
        if fmt in ("all", "pdf"):
            results["pdf"] = export_to_pdf(html_file, 
                str(Path(out_dir) / f"AI_Weekly_Report_{datetime.now().strftime('%Y%m%d')}.pdf"))
        if fmt in ("all", "docx"):
            results["docx"] = export_to_docx(html_file,
                str(Path(out_dir) / f"AI_Weekly_Report_{datetime.now().strftime('%Y%m%d')}.docx"))
        if fmt in ("all", "json"):
            results["json"] = export_json_report(html_file,
                str(Path(out_dir) / f"AI_Weekly_Data_{datetime.now().strftime('%Y%m%d')}.json"))
        return results

    if args.watch:
        print(f"👀 监听模式: {html_file}")
        print(f"   文件变更时自动导出 (Ctrl+C 停止)\n")
        last_mtime = Path(html_file).stat().st_mtime
        try:
            while True:
                time.sleep(5)
                mtime = Path(html_file).stat().st_mtime
                if mtime != last_mtime:
                    last_mtime = mtime
                    print(f"\n🔄 检测到文件变更，重新导出...")
                    do_export()
        except KeyboardInterrupt:
            print("\n👋 监听已停止")
    else:
        results = do_export()
        print(f"\n📁 导出完成: {len([v for v in results.values() if v])} 个文件")


if __name__ == "__main__":
    main()
